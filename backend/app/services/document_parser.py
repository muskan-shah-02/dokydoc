# DAE-01/02 FIX: Robust Document Parser with Multiple Fallback Strategies
# This module implements a resilient PDF parsing system with OCR support

import os
import tempfile
from pathlib import Path
from typing import Tuple, List, Optional
from enum import Enum
import google.generativeai as genai
from tenacity import retry, stop_after_attempt, wait_exponential
import fitz  # PyMuPDF for PDF image extraction
from docx import Document as DocxDocument
from PIL import Image
import io
import asyncio

from app.core.config import settings
from app.core.logging import LoggerMixin
from app.services.ai.prompt_manager import prompt_manager, PromptType


class ParserStrategy(Enum):
    """DAE-01/02: Enumeration of available PDF parsing strategies."""
    PYMUPDF = "pymupdf"  # Primary: Fast and reliable
    PDFPLUMBER = "pdfplumber"  # Fallback 1: Better text extraction
    PYPDF2 = "pypdf2"  # Fallback 2: Basic fallback
    OCR = "ocr"  # Fallback 3: For scanned PDFs


class MultiModalDocumentParser(LoggerMixin):
    """
    DAE-01/02 FIX: A robust service to parse text content from documents using multiple strategies.

    Features:
    - Multiple PDF parsing strategies (PyMuPDF, pdfplumber, PyPDF2)
    - OCR support for scanned PDFs (pytesseract)
    - Automatic fallback on parser failures
    - Detailed error messages for debugging
    """

    def __init__(self):
        super().__init__()

        if not settings.GEMINI_API_KEY or "YOUR_GEMINI_API_KEY_HERE" in settings.GEMINI_API_KEY:
            raise ValueError("GEMINI_API_KEY is not configured correctly in app/core/config.py.")

        genai.configure(api_key=settings.GEMINI_API_KEY)
        self.model = genai.GenerativeModel(settings.GEMINI_MODEL)
        self.vision_model = genai.GenerativeModel(settings.GEMINI_VISION_MODEL)
        self.logger.info("MultiModalDocumentParser initialized with fallback support.")

        # Track which parser succeeded (for logging/analytics)
        self.successful_strategy: Optional[ParserStrategy] = None

    def _get_file_extension(self, file_path: str) -> str:
        return os.path.splitext(file_path)[1].lower()

    def _is_supported_directly(self, file_path: str) -> bool:
        extension = self._get_file_extension(file_path)
        return extension in SUPPORTED_MIME_TYPES

    def _requires_conversion(self, file_path: str) -> bool:
        extension = self._get_file_extension(file_path)
        return extension in CONVERSION_REQUIRED

    # ==================== DAE-01/02: PDF PARSING STRATEGIES ====================

    def _extract_text_pymupdf(self, file_path: str) -> Tuple[str, bool]:
        """
        Strategy 1: PyMuPDF (fitz) - Fast and reliable.

        Returns:
            Tuple[str, bool]: (extracted_text, is_scanned)
        """
        try:
            doc = fitz.open(file_path)
            text_parts = []
            is_scanned = True  # Assume scanned until proven otherwise

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                text_parts.append(text)

                # If any page has text, it's not purely scanned
                if text.strip():
                    is_scanned = False

            doc.close()
            combined_text = "\n\n".join(text_parts)

            if not combined_text.strip():
                self.logger.warning("PyMuPDF extracted no text - possibly scanned PDF")
                return "", True

            self.logger.info(f"✅ PyMuPDF extracted {len(combined_text)} characters")
            return combined_text, is_scanned

        except Exception as e:
            self.logger.error(f"❌ PyMuPDF failed: {e}")
            raise

    def _detect_columns(self, page) -> int:
        """
        CAE-03 FIX: Detect number of text columns on a pdfplumber page
        using word x-coordinate clustering.

        Returns:
            int: Number of detected columns (1 if single-column or uncertain)
        """
        try:
            words = page.extract_words(x_tolerance=3, y_tolerance=3)
            if len(words) < 10:
                return 1

            page_width = page.width
            midpoint = page_width / 2
            gap_threshold = page_width * 0.08  # 8% gap around midpoint

            left_words = [w for w in words if float(w["x1"]) < midpoint - gap_threshold]
            right_words = [w for w in words if float(w["x0"]) > midpoint + gap_threshold]

            # Multi-column if both sides have substantial content
            if len(left_words) > 5 and len(right_words) > 5:
                left_ratio = len(left_words) / len(words)
                right_ratio = len(right_words) / len(words)
                if left_ratio > 0.2 and right_ratio > 0.2:
                    self.logger.info(
                        f"CAE-03: Detected 2-column layout "
                        f"(left={len(left_words)}, right={len(right_words)}, total={len(words)})"
                    )
                    return 2

            return 1
        except Exception:
            return 1

    def _extract_columns_merged(self, page) -> str:
        """
        CAE-03 FIX: Extract text from a multi-column page by reading
        left column first, then right column (top-to-bottom for each).
        """
        try:
            page_width = page.width
            midpoint = page_width / 2

            # Crop left and right halves with small overlap tolerance
            left_bbox = (0, 0, midpoint + 2, page.height)
            right_bbox = (midpoint - 2, 0, page_width, page.height)

            left_text = page.within_bbox(left_bbox).extract_text() or ""
            right_text = page.within_bbox(right_bbox).extract_text() or ""

            # Combine: left column first, then right column
            combined = left_text.strip()
            if right_text.strip():
                combined += "\n\n" + right_text.strip()

            return combined
        except Exception as e:
            self.logger.warning(f"CAE-03: Column extraction failed, falling back: {e}")
            return page.extract_text() or ""

    def _extract_text_pdfplumber(self, file_path: str) -> Tuple[str, bool]:
        """
        Strategy 2: pdfplumber - Better text extraction for complex layouts.

        CAE-03 FIX: Now detects multi-column layouts and extracts columns
        in reading order (left-to-right, top-to-bottom per column).

        Returns:
            Tuple[str, bool]: (extracted_text, is_scanned)
        """
        try:
            import pdfplumber

            text_parts = []
            is_scanned = True
            multi_column_pages = 0

            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                for page in pdf.pages:
                    # CAE-03: Detect column layout per page
                    num_columns = self._detect_columns(page)

                    if num_columns >= 2:
                        text = self._extract_columns_merged(page)
                        multi_column_pages += 1
                    else:
                        text = page.extract_text()

                    if text:
                        text_parts.append(text)
                        is_scanned = False

            combined_text = "\n\n".join(text_parts)

            if not combined_text.strip():
                self.logger.warning("pdfplumber extracted no text - possibly scanned PDF")
                return "", True

            if multi_column_pages > 0:
                self.logger.info(
                    f"CAE-03: {multi_column_pages}/{total_pages} pages had multi-column layout"
                )

            self.logger.info(f"✅ pdfplumber extracted {len(combined_text)} characters")
            return combined_text, is_scanned

        except Exception as e:
            self.logger.error(f"❌ pdfplumber failed: {e}")
            raise

    def _extract_text_pypdf2(self, file_path: str) -> Tuple[str, bool]:
        """
        Strategy 3: PyPDF2 - Basic fallback parser.

        Returns:
            Tuple[str, bool]: (extracted_text, is_scanned)
        """
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []
            is_scanned = True

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
                    is_scanned = False

            combined_text = "\n\n".join(text_parts)

            if not combined_text.strip():
                self.logger.warning("PyPDF2 extracted no text - possibly scanned PDF")
                return "", True

            self.logger.info(f"✅ PyPDF2 extracted {len(combined_text)} characters")
            return combined_text, is_scanned

        except Exception as e:
            self.logger.error(f"❌ PyPDF2 failed: {e}")
            raise

    def _extract_text_ocr(self, file_path: str) -> Tuple[str, bool]:
        """
        Strategy 4: OCR with pytesseract - For scanned PDFs.

        Returns:
            Tuple[str, bool]: (extracted_text, is_scanned)
        """
        try:
            import pytesseract
            from pdf2image import convert_from_path

            self.logger.info("🔍 Using OCR for scanned PDF...")

            # Convert PDF to images
            images = convert_from_path(file_path)
            text_parts = []

            for i, image in enumerate(images):
                self.logger.info(f"OCR processing page {i+1}/{len(images)}")
                text = pytesseract.image_to_string(image)
                if text.strip():
                    text_parts.append(text)

            combined_text = "\n\n".join(text_parts)

            if not combined_text.strip():
                self.logger.warning("OCR extracted no text - PDF may be empty or corrupted")
                return "", True

            self.logger.info(f"✅ OCR extracted {len(combined_text)} characters from {len(images)} pages")
            return combined_text, True  # OCR is always for scanned PDFs

        except ImportError as e:
            self.logger.error(f"❌ OCR dependencies missing: {e}. Install: pip install pytesseract pdf2image")
            raise ValueError("OCR dependencies not installed. Cannot process scanned PDF.")
        except Exception as e:
            self.logger.error(f"❌ OCR failed: {e}")
            raise

    def _parse_pdf_with_fallbacks(self, file_path: str) -> str:
        """
        DAE-01/02 FIX: Parse PDF using multiple strategies with automatic fallback.

        Strategy Order:
        1. PyMuPDF (fast, reliable)
        2. pdfplumber (better layout handling)
        3. PyPDF2 (basic fallback)
        4. OCR (for scanned PDFs)

        Returns:
            str: Extracted text content

        Raises:
            ValueError: If all strategies fail
        """
        errors = []
        strategies = [
            (ParserStrategy.PYMUPDF, self._extract_text_pymupdf),
            (ParserStrategy.PDFPLUMBER, self._extract_text_pdfplumber),
            (ParserStrategy.PYPDF2, self._extract_text_pypdf2),
        ]

        for strategy, extractor in strategies:
            try:
                self.logger.info(f"📄 Trying {strategy.value} parser...")
                text, is_scanned = extractor(file_path)

                if text.strip():
                    self.successful_strategy = strategy
                    return text

                if is_scanned:
                    # PDF appears to be scanned, try OCR
                    self.logger.info("PDF appears to be scanned, falling back to OCR")
                    break

            except Exception as e:
                error_msg = f"{strategy.value}: {str(e)}"
                errors.append(error_msg)
                self.logger.warning(f"⚠️ {strategy.value} failed, trying next strategy...")
                continue

        # If we get here, either all strategies failed or PDF is scanned
        # Try OCR as last resort
        try:
            self.logger.info("📄 Trying OCR parser (last resort)...")
            text, _ = self._extract_text_ocr(file_path)
            if text.strip():
                self.successful_strategy = ParserStrategy.OCR
                return text
        except Exception as e:
            error_msg = f"OCR: {str(e)}"
            errors.append(error_msg)

        # All strategies failed
        error_summary = "\n".join([f"  - {err}" for err in errors])
        raise ValueError(
            f"All PDF parsing strategies failed for {os.path.basename(file_path)}:\n{error_summary}\n\n"
            f"Possible issues:\n"
            f"  - PDF may be corrupted or encrypted\n"
            f"  - File may not be a valid PDF\n"
            f"  - OCR dependencies may be missing (for scanned PDFs)"
        )

    # ==================== ORIGINAL METHODS (UNCHANGED) ====================

    def _extract_images_from_pdf(self, file_path: str) -> list:
        images = []
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n - pix.alpha < 4:
                            img_data = pix.tobytes("png")
                        else:
                            pix1 = fitz.Pixmap(fitz.csRGB, pix)
                            img_data = pix1.tobytes("png")
                            pix1 = None
                        images.append({'page': page_num + 1, 'index': img_index, 'data': img_data, 'format': 'png'})
                        pix = None
                    except Exception:
                        continue
            doc.close()
            return images
        except Exception as e:
            self.logger.error(f"Error extracting images from PDF: {e}")
            return []

    def _extract_images_from_docx(self, file_path: str) -> list:
        images = []
        try:
            doc = DocxDocument(file_path)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        images.append({'data': rel.target_part.blob, 'format': 'png', 'width': None, 'height': None})
                    except Exception:
                        continue
            return images
        except Exception as e:
            self.logger.error(f"Error extracting images from DOCX: {e}")
            return []

    async def _analyze_image_with_vision(self, image_data: bytes, image_index: int) -> str:
        """
        Analyze an image using Gemini Vision API.
        Wraps the synchronous call in a thread to be async-compatible.
        """
        try:
            prompt = prompt_manager.get_prompt(PromptType.IMAGE_ANALYSIS)
            image = Image.open(io.BytesIO(image_data))

            # FIX: Run synchronous SDK call in a thread
            response = await asyncio.to_thread(
                self.vision_model.generate_content, [prompt, image]
            )

            return f"[image-{image_index:02d}: {response.text}]"
        except Exception as e:
            self.logger.error(f"Error analyzing image {image_index}: {e}")
            return f"[image-{image_index:02d}: Image analysis failed]"

    def _convert_docx_to_text(self, file_path: str) -> str:
        try:
            doc = DocxDocument(file_path)
            text_content = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            return "\n".join(text_content)
        except Exception as e:
            raise ValueError(f"Failed to convert DOCX file: {e}")

    def _convert_doc_to_text(self, file_path: str) -> str:
        try:
            import docx2txt
            return docx2txt.process(file_path).strip()
        except Exception as e:
            raise ValueError(f"Failed to convert DOC file: {e}")

    @retry(
        stop=stop_after_attempt(6),
        wait=wait_exponential(multiplier=2, min=4, max=60)
    )
    async def _process_with_gemini(self, file_path: str, mime_type: str, prompt: str) -> tuple[str, int, int]:
        """
        Process a file with Gemini API with retry logic.
        FIXED: Uses asyncio.to_thread to handle synchronous SDK calls.

        Returns:
            Tuple[str, int, int]: (text, input_tokens, output_tokens)
        """
        try:
            with open(file_path, 'rb') as f:
                file_data = f.read()

            # Prepare arguments for the sync call
            args = []
            if mime_type.startswith("text/"):
                try:
                    text_content = file_data.decode("utf-8")
                except UnicodeDecodeError:
                    text_content = file_data.decode("latin-1", errors="replace")
                args = [prompt, text_content]
            else:
                doc_blob = {"mime_type": mime_type, "data": file_data}
                args = [prompt, doc_blob]

            # --- CRITICAL FIX START ---
            # Run the blocking synchronous Google call in a separate thread
            # This allows 'await' to work correctly without crashing the worker
            response = await asyncio.to_thread(self.model.generate_content, args)
            # --- CRITICAL FIX END ---

            # Extract token counts from usage_metadata
            input_tokens = 0
            output_tokens = 0
            if hasattr(response, 'usage_metadata'):
                input_tokens = getattr(response.usage_metadata, 'prompt_token_count', 0)
                output_tokens = getattr(response.usage_metadata, 'candidates_token_count', 0)

                self.logger.info(
                    f"📊 Token usage: {input_tokens} input + {output_tokens} output = "
                    f"{input_tokens + output_tokens} total tokens"
                )

            return response.text, input_tokens, output_tokens

        except Exception as e:
            self.logger.error(f"Error processing file {file_path} with Gemini: {e}")
            raise

    async def parse_with_images(self, file_path: str) -> tuple[str, int, int]:
        """
        Main orchestration method.

        DAE-01/02 FIX: Now uses fallback strategies for robust PDF parsing.

        Returns:
            Tuple[str, int, int]: (text, input_tokens, output_tokens)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_path = str(file_path)
        extension = self._get_file_extension(file_path)
        self.logger.info(f"Starting parsing for: {file_path}")

        try:
            if self._requires_conversion(file_path):
                self.logger.info(f"Converting {extension} file to text")
                if extension == ".docx":
                    # Conversion doesn't use Gemini, so 0 tokens
                    return self._convert_docx_to_text(file_path), 0, 0
                elif extension == ".doc":
                    return self._convert_doc_to_text(file_path), 0, 0

            elif self._is_supported_directly(file_path):
                mime_type = SUPPORTED_MIME_TYPES[extension]

                if extension == ".pdf":
                    # DAE-01/02 FIX: Use fallback strategies for PDF parsing
                    pdf_text = self._parse_pdf_with_fallbacks(file_path)

                    # Extract images for additional context
                    images = self._extract_images_from_pdf(file_path)
                    image_descriptions = []
                    vision_input_tokens = 0
                    vision_output_tokens = 0

                    if images:
                        self.logger.info(f"Found {len(images)} images in PDF.")
                        for i, img in enumerate(images):
                            desc = await self._analyze_image_with_vision(img['data'], i)
                            image_descriptions.append(desc)
                            # Note: Vision API token tracking would need separate implementation

                    if image_descriptions:
                        pdf_text += "\n\n" + "\n".join(image_descriptions)

                    self.logger.info(
                        f"✅ PDF parsed successfully using {self.successful_strategy.value if self.successful_strategy else 'unknown'}"
                    )

                    # Note: Fallback parsers don't use Gemini, so 0 tokens
                    return pdf_text, 0, 0

                return await self._process_with_gemini(file_path, mime_type, "Extract all text content.")

            else:
                raise ValueError(f"Unsupported file type: {extension}")

        except Exception as e:
            self.logger.error(f"Error parsing document: {e}")
            raise


SUPPORTED_MIME_TYPES = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/plain",
    ".html": "text/html",
    ".xml": "application/xml",
}

CONVERSION_REQUIRED = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
}
