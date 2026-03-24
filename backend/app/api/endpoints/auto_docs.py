"""
Auto Docs API Endpoints
Sprint 8: AI-powered documentation generation (Module 12).
Sprint 9: Multi-source generation.

  POST /auto-docs/generate        — Generate from a single source (document or repository)
  POST /auto-docs/generate-multi  — Generate from multiple combined sources
  GET  /auto-docs/                — List generated docs for current tenant
  GET  /auto-docs/{id}            — Get a single generated doc
  GET  /auto-docs/{id}/export     — Export as .docx or .pdf
  POST /auto-docs/{id}/refine     — Refine the document with an AI prompt
"""
import io
import re
import asyncio
from typing import Any, Optional
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response as FastAPIResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from app import models
from app.api import deps
from app.db.session import get_db
from app.crud.crud_generated_doc import crud_generated_doc
from app.services.auto_docs_service import auto_docs_service, _DOC_TYPE_TITLES
from app.core.logging import get_logger

logger = get_logger("api.auto_docs")

router = APIRouter()


# ---- Schemas ----

class GenerateDocRequest(BaseModel):
    source_type: str = Field(..., pattern="^(document|repository)$")
    source_id: int
    doc_type: str = Field(..., description=(
        "component_spec | architecture_diagram | api_summary | brd | test_cases | data_models"
    ))


class SourceEntry(BaseModel):
    type: str = Field(
        ...,
        pattern="^(document|repository|code_file|standalone|jira_item|analysis|folder|api_endpoint|confluence_page)$",
    )
    id: int
    folder_path: Optional[str] = Field(None, description="For type='folder': path prefix within the repo")
    endpoint_path: Optional[str] = Field(None, description="For type='api_endpoint': e.g. 'POST /api/v1/generate'")


class GenerateMultiRequest(BaseModel):
    sources: list[SourceEntry] = Field(..., min_length=1, description="One or more sources to combine")
    doc_type: str = Field(..., description=(
        "component_spec | architecture_diagram | api_summary | brd | test_cases | data_models | "
        "class_diagram | sequence_diagram | code_flow_diagram | component_interaction_diagram | "
        "folder_architecture | api_data_flow"
    ))


class RecommendRequest(BaseModel):
    sources: list[SourceEntry] = Field(..., min_length=1)


class RecommendResponse(BaseModel):
    recommended: str
    reason: str
    alternatives: list[str]
    confidence: float


class RefineRequest(BaseModel):
    prompt: str = Field(..., min_length=1, description="Instruction for how to refine/update the document")
    save_as_new: bool = Field(False, description="If true, save the refined version as a new generated doc")


class GeneratedDocResponse(BaseModel):
    id: int
    tenant_id: int
    user_id: Optional[int] = None
    source_type: str
    source_id: int
    source_name: Optional[str] = None
    doc_type: str
    title: str
    content: str
    metadata: Optional[dict] = None
    status: str
    created_at: str

    class Config:
        from_attributes = True


# ---- Export helpers ----

def _slugify(text: str) -> str:
    """Make a safe filename from a title."""
    s = re.sub(r"[^\w\s-]", "", text.lower())
    s = re.sub(r"[\s_-]+", "-", s).strip("-")
    return s[:80] or "document"


def _markdown_to_docx(title: str, content: str) -> io.BytesIO:
    """Convert markdown text to a .docx file using python-docx."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown line by line
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Pt(18)
            i += 1
            continue

        # ATX headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)), 4)
            text = m.group(2).strip()
            doc.add_heading(_strip_inline_md(text), level=level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}\s*$", line):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # Unordered list
        if re.match(r"^\s*[-*+]\s+", line):
            text = re.sub(r"^\s*[-*+]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, text)
            i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, text)
            i += 1
            continue

        # Table row
        if line.startswith("|"):
            # Collect all table rows
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row_line = lines[i].strip("|").strip()
                cells = [c.strip() for c in row_line.split("|")]
                if not all(re.match(r"^[-:]+$", c) for c in cells if c):
                    table_rows.append(cells)
                i += 1
            if table_rows:
                col_count = max(len(r) for r in table_rows)
                t = doc.add_table(rows=len(table_rows), cols=col_count)
                t.style = "Table Grid"
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = t.cell(r_idx, c_idx)
                        cell.text = _strip_inline_md(cell_text)
                        if r_idx == 0:
                            cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else None
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, line)
        i += 1

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _strip_inline_md(text: str) -> str:
    """Strip bold/italic/code markers from text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _add_inline_runs(paragraph: Any, text: str) -> None:
    """Add runs with bold/italic/code formatting to a paragraph."""
    # Split on bold (**text**), italic (*text*), inline code (`text`)
    pattern = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def _markdown_to_pdf(title: str, content: str) -> io.BytesIO:
    """Convert markdown text to a simple PDF using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    doc = fitz.open()
    PAGE_W, PAGE_H = 595, 842  # A4 in points
    MARGIN = 60
    TEXT_W = PAGE_W - 2 * MARGIN
    LINE_H = 14
    PARA_GAP = 6

    def new_page():
        page = doc.new_page(width=PAGE_W, height=PAGE_H)
        return page, MARGIN  # page, current_y

    def write_line(page, y, text, fontsize=11, bold=False, color=(0, 0, 0), indent=0):
        font = "helv" if not bold else "hebo"
        page.insert_text(
            (MARGIN + indent, y),
            text[:200],  # safety cap per line
            fontname=font,
            fontsize=fontsize,
            color=color,
        )
        return y + fontsize + 3

    def wrap_text(text, max_chars=90):
        """Very simple word wrapper."""
        words = text.split()
        lines = []
        current = ""
        for w in words:
            if len(current) + len(w) + 1 <= max_chars:
                current = (current + " " + w).strip()
            else:
                if current:
                    lines.append(current)
                current = w
        if current:
            lines.append(current)
        return lines or [""]

    page, y = new_page()

    def ensure_space(needed):
        nonlocal page, y
        if y + needed > PAGE_H - MARGIN:
            page, y = new_page()

    # Title
    ensure_space(40)
    y = write_line(page, y, title, fontsize=18, bold=True, color=(0.13, 0.27, 0.53))
    y += 10

    # Parse markdown
    lines_md = content.split("\n")
    i = 0
    while i < len(lines_md):
        line = lines_md[i]

        # Code block
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines_md) and not lines_md[i].strip().startswith("```"):
                code_lines.append(lines_md[i])
                i += 1
            ensure_space(LINE_H * min(len(code_lines), 5) + PARA_GAP)
            for cl in code_lines[:30]:
                ensure_space(LINE_H)
                y = write_line(page, y, cl[:100], fontsize=8, color=(0.3, 0.3, 0.3), indent=20)
            y += PARA_GAP
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = _strip_inline_md(m.group(2))
            fontsize = max(10, 18 - (level - 1) * 3)
            ensure_space(fontsize + 10)
            y += PARA_GAP
            y = write_line(page, y, text, fontsize=fontsize, bold=True, color=(0.13, 0.27, 0.53))
            y += PARA_GAP
            i += 1
            continue

        # HR
        if re.match(r"^[-*_]{3,}\s*$", line):
            ensure_space(12)
            page.draw_line((MARGIN, y), (PAGE_W - MARGIN, y), color=(0.8, 0.8, 0.8))
            y += 10
            i += 1
            continue

        # List item
        if re.match(r"^\s*[-*+]\s+", line):
            text = re.sub(r"^\s*[-*+]\s+", "• ", line)
            text = _strip_inline_md(text)
            for wl in wrap_text(text):
                ensure_space(LINE_H)
                y = write_line(page, y, wl, fontsize=10, indent=15)
            y += 2
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            m2 = re.match(r"^\s*(\d+)\.\s+(.*)", line)
            if m2:
                text = f"{m2.group(1)}. {_strip_inline_md(m2.group(2))}"
                for wl in wrap_text(text):
                    ensure_space(LINE_H)
                    y = write_line(page, y, wl, fontsize=10, indent=15)
                y += 2
            i += 1
            continue

        # Table row (simplified)
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            is_sep = all(re.match(r"^[-:]+$", c) for c in cells if c)
            if not is_sep:
                row_text = "  |  ".join(_strip_inline_md(c) for c in cells if c)
                ensure_space(LINE_H)
                y = write_line(page, y, row_text[:120], fontsize=9)
            i += 1
            continue

        # Blank
        if not line.strip():
            y += PARA_GAP
            i += 1
            continue

        # Normal paragraph
        text = _strip_inline_md(line)
        for wl in wrap_text(text):
            ensure_space(LINE_H)
            y = write_line(page, y, wl, fontsize=10)
        y += PARA_GAP
        i += 1

    bio = io.BytesIO(doc.tobytes())
    bio.seek(0)
    return bio


# ---- Endpoints ----

@router.post("/generate", status_code=201)
async def generate_doc(
    payload: GenerateDocRequest,
    tenant_id: int = Depends(deps.get_tenant_id),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    """
    Generate a documentation artifact using AI.
    This is a synchronous call — the response contains the generated content.
    """
    if payload.doc_type not in auto_docs_service.SUPPORTED_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown doc_type. Supported: {auto_docs_service.SUPPORTED_TYPES}",
        )

    try:
        result = await auto_docs_service.generate(
            db,
            doc_type=payload.doc_type,
            source_type=payload.source_type,
            source_id=payload.source_id,
            tenant_id=tenant_id,
            user_id=current_user.id,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    obj = crud_generated_doc.create(
        db,
        tenant_id=tenant_id,
        user_id=current_user.id,
        source_type=payload.source_type,
        source_id=payload.source_id,
        source_name=result.get("source_name"),
        doc_type=payload.doc_type,
        title=result["title"],
        content=result["content"],
        metadata=result.get("metadata"),
        status=result["status"],
    )

    return {
        "id": obj.id,
        "tenant_id": obj.tenant_id,
        "user_id": obj.user_id,
        "source_type": obj.source_type,
        "source_id": obj.source_id,
        "source_name": obj.source_name,
        "doc_type": obj.doc_type,
        "title": obj.title,
        "content": obj.content,
        "metadata": obj.doc_metadata,
        "status": obj.status,
        "created_at": obj.created_at.isoformat(),
    }


@router.post("/generate-multi", status_code=201)
async def generate_doc_multi(
    payload: GenerateMultiRequest,
    tenant_id: int = Depends(deps.get_tenant_id),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    """
    Generate a documentation artifact from multiple combined sources.
    Accepts 1–N documents and/or repositories as inputs; the AI merges all context.
    """
    if payload.doc_type not in auto_docs_service.SUPPORTED_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown doc_type. Supported: {auto_docs_service.SUPPORTED_TYPES}",
        )

    sources = [
        {
            "type": s.type,
            "id": s.id,
            "folder_path": s.folder_path or "",
            "endpoint_path": s.endpoint_path or "",
        }
        for s in payload.sources
    ]

    try:
        result = await auto_docs_service.generate_multi(
            db,
            doc_type=payload.doc_type,
            sources=sources,
            tenant_id=tenant_id,
            user_id=current_user.id,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # If only one source, use its type/id for backward compat; otherwise "multi"/0
    if len(sources) == 1:
        src_type = sources[0]["type"]
        src_id = sources[0]["id"]
    else:
        src_type = "multi"
        src_id = 0

    obj = crud_generated_doc.create(
        db,
        tenant_id=tenant_id,
        user_id=current_user.id,
        source_type=src_type,
        source_id=src_id,
        source_name=result.get("source_name"),
        doc_type=payload.doc_type,
        title=result["title"],
        content=result["content"],
        metadata=result.get("metadata"),
        status=result["status"],
        source_ids=result.get("source_ids"),
    )

    return {
        "id": obj.id,
        "tenant_id": obj.tenant_id,
        "user_id": obj.user_id,
        "source_type": obj.source_type,
        "source_id": obj.source_id,
        "source_name": obj.source_name,
        "source_ids": obj.source_ids,
        "doc_type": obj.doc_type,
        "title": obj.title,
        "content": obj.content,
        "metadata": obj.doc_metadata,
        "status": obj.status,
        "created_at": obj.created_at.isoformat(),
    }


@router.get("/")
def list_generated_docs(
    tenant_id: int = Depends(deps.get_tenant_id),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user),
    source_type: Optional[str] = Query(None),
    source_id: Optional[int] = Query(None),
    doc_type: Optional[str] = Query(None),
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
) -> Any:
    """List previously generated docs with optional filters."""
    docs = crud_generated_doc.list_for_tenant(
        db,
        tenant_id=tenant_id,
        source_type=source_type,
        source_id=source_id,
        doc_type=doc_type,
        skip=skip,
        limit=limit,
    )
    return {
        "docs": [
            {
                "id": d.id,
                "source_type": d.source_type,
                "source_id": d.source_id,
                "source_name": d.source_name,
                "doc_type": d.doc_type,
                "title": d.title,
                "status": d.status,
                "created_at": d.created_at.isoformat(),
            }
            for d in docs
        ],
        "total": len(docs),
        "supported_types": auto_docs_service.SUPPORTED_TYPES,
    }


@router.get("/{doc_id}/export")
def export_generated_doc(
    doc_id: int,
    format: str = Query(..., pattern="^(docx|pdf)$", description="Export format: docx or pdf"),
    tenant_id: int = Depends(deps.get_tenant_id),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    """Export a generated doc as Word (.docx) or PDF (.pdf)."""
    obj = crud_generated_doc.get_by_id(db, doc_id=doc_id, tenant_id=tenant_id)
    if not obj:
        raise HTTPException(status_code=404, detail="Generated doc not found")

    content = obj.content or ""
    title = obj.title or f"Document {doc_id}"
    slug = _slugify(title)

    try:
        if format == "docx":
            bio = _markdown_to_docx(title, content)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{slug}.docx"
        else:
            bio = _markdown_to_pdf(title, content)
            media_type = "application/pdf"
            filename = f"{slug}.pdf"
    except Exception as e:
        logger.error(f"Export failed for doc {doc_id} as {format}: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e}")

    return FastAPIResponse(
        content=bio.getvalue(),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/{doc_id}/refine", status_code=200)
async def refine_generated_doc(
    doc_id: int,
    payload: RefineRequest,
    tenant_id: int = Depends(deps.get_tenant_id),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    """
    Refine a generated document using an AI prompt.
    The current document content is sent to the AI along with the user's refinement
    instruction. The refined result is returned and optionally saved as a new doc.
    """
    obj = crud_generated_doc.get_by_id(db, doc_id=doc_id, tenant_id=tenant_id)
    if not obj:
        raise HTTPException(status_code=404, detail="Generated doc not found")

    refinement_prompt = f"""You are a technical documentation editor. Below is an existing document followed by a refinement request from the user.

Apply the requested changes and return the **complete refined document** in Markdown format. Maintain the professional structure and quality of the document. Do not truncate — return the full document.

--- EXISTING DOCUMENT ---
{(obj.content or "")[:10000]}
--- END OF DOCUMENT ---

User's refinement instruction:
{payload.prompt}

Return only the refined Markdown document, nothing else."""

    try:
        from app.services.ai.gemini import gemini_service

        response = await gemini_service.generate_content(
            refinement_prompt,
            tenant_id=tenant_id,
            user_id=current_user.id,
            operation="auto_docs_refine",
        )
        refined_content = response.text or ""
        tokens = gemini_service.extract_token_usage(response)
    except Exception as e:
        logger.error(f"AutoDocs refinement failed for doc {doc_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Refinement failed: {e}")

    if payload.save_as_new:
        new_obj = crud_generated_doc.create(
            db,
            tenant_id=tenant_id,
            user_id=current_user.id,
            source_type=obj.source_type,
            source_id=obj.source_id,
            source_name=obj.source_name,
            doc_type=obj.doc_type,
            title=f"{obj.title} (refined)",
            content=refined_content,
            metadata={
                "doc_type": obj.doc_type,
                "refined_from": doc_id,
                "refine_prompt": payload.prompt[:200],
                **tokens,
            },
            status="completed",
            source_ids=obj.source_ids,
        )
        return {
            "content": refined_content,
            "saved": True,
            "id": new_obj.id,
            "title": new_obj.title,
            "created_at": new_obj.created_at.isoformat(),
        }

    return {
        "content": refined_content,
        "saved": False,
        "id": None,
    }


@router.get("/{doc_id}")
def get_generated_doc(
    doc_id: int,
    tenant_id: int = Depends(deps.get_tenant_id),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    """Get full content of a specific generated doc."""
    obj = crud_generated_doc.get_by_id(db, doc_id=doc_id, tenant_id=tenant_id)
    if not obj:
        raise HTTPException(status_code=404, detail="Generated doc not found")
    return {
        "id": obj.id,
        "tenant_id": obj.tenant_id,
        "user_id": obj.user_id,
        "source_type": obj.source_type,
        "source_id": obj.source_id,
        "source_name": obj.source_name,
        "source_ids": obj.source_ids,
        "doc_type": obj.doc_type,
        "title": obj.title,
        "content": obj.content,
        "metadata": obj.doc_metadata,
        "status": obj.status,
        "created_at": obj.created_at.isoformat(),
    }

@router.post("/recommend", status_code=200)
def recommend_diagram_type(
    payload: RecommendRequest,
    tenant_id: int = Depends(deps.get_tenant_id),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    """
    Recommend the best diagram/doc type for the given sources.
    Analyzes structured_analysis from each source to determine the most
    appropriate visualization (class_diagram, sequence_diagram, etc.).
    """
    sources = [
        {
            "type": s.type,
            "id": s.id,
            "folder_path": s.folder_path or "",
            "endpoint_path": s.endpoint_path or "",
        }
        for s in payload.sources
    ]
    try:
        result = auto_docs_service.recommend_diagram_type(db, sources, tenant_id)
    except Exception as e:
        logger.warning(f"Recommendation failed: {e}")
        result = {
            "recommended": "component_spec",
            "reason": "Could not analyze sources — defaulting to Component Specification.",
            "alternatives": ["architecture_diagram", "api_summary"],
            "confidence": 0.5,
        }
    return result


@router.get("/discovered-endpoints")
def get_discovered_endpoints(
    repo_id: Optional[int] = Query(None, description="Filter to a specific repository"),
    q: str = Query("", description="Search query — filters by path or method"),
    tenant_id: int = Depends(deps.get_tenant_id),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    """
    Return all API endpoints discovered during code analysis.
    Searches all analyzed components' api_contracts.
    Supports optional filtering by repo and search query.
    """
    endpoints = auto_docs_service.get_discovered_endpoints(
        db, tenant_id=tenant_id, repo_id=repo_id, query=q
    )
    return {"endpoints": endpoints, "total": len(endpoints)}
